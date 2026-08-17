from docx import Document
import os

# Paths
here = os.path.dirname(__file__)
md_path = os.path.join(here, 'chapter3.md')
out_path = os.path.join(here, 'Chapter_3_System_Design_and_Implementation.docx')

# Read markdown-like file and convert to simple Word doc
with open(md_path, 'r', encoding='utf-8') as f:
    text = f.read()

# Create document
doc = Document()

# Simple parsing: lines starting with '#', '##', '###' become headings
for block in text.split('\n\n'):
    line = block.strip()
    if not line:
        continue
    if line.startswith('Chapter 3'):
        doc.add_heading(line, level=1)
        continue
    if line.startswith('3.') or line.startswith('3'):
        # treat section headers that start with numbers as heading
        first_line = line.split('\n', 1)[0]
        doc.add_heading(first_line, level=2)
        rest = line[len(first_line):].strip()
        if rest:
            doc.add_paragraph(rest)
        continue
    # default: add as paragraph (preserve newlines inside the block)
    doc.add_paragraph(line)

# Save
doc.save(out_path)
print('Saved:', out_path)
